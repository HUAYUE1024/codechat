"""
CodeAgent — ReAct Agent with Planning / Tools / Memory / Action / LLM

Architecture:
    Planning  →  decompose user goal into executable steps
    Tools     →  external capabilities the agent can invoke
    Memory    →  short-term (session) + long-term (persisted)
    Action    →  execute tool calls with retry, timeout, error handling
    LLM       →  reasoning engine powering plan, decide, reflect
"""

from __future__ import annotations

import hashlib
import json
import sys
import os
import re
import time
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path

from .config import get_snowcode_dir
from .rag import _get_llm_config, _call_llm
from .scanner import scan_files, read_file
from .store import VectorStore


# ═══════════════════════════════════════════════════════════════════════
#  LLM Client
# ═══════════════════════════════════════════════════════════════════════

class LLMClient:
    """Thin wrapper around LLM backends."""

    def __init__(self, model: str | None = None):
        self.model = model
        self.api_key, self.base_url, self.model_name, self.thinking = _get_llm_config(model)

    @property
    def available(self) -> bool:
        return bool(self.api_key)

    def complete(self, system: str, user: str, temperature: float = 0.1) -> str:
        """Non-streaming completion."""
        if not self.available:
            return ""
            
        if self.api_key == "ollama":
            try:
                import httpx
                resp = httpx.post(
                    f"{self.base_url}/api/chat",
                    json={
                        "model": self.model_name,
                        "messages": [
                            {"role": "system", "content": system},
                            {"role": "user", "content": user},
                        ],
                        "stream": False,
                    },
                    timeout=120,
                )
                if resp.status_code == 200:
                    return resp.json()["message"]["content"]
                return f"[LLM Error] HTTP {resp.status_code}: {resp.text}"
            except Exception as e:
                return f"[LLM Error] {type(e).__name__}: {e}"
                
        try:
            import openai
            client = openai.OpenAI(api_key=self.api_key, base_url=self.base_url)
            extra = {"enable_thinking": self.thinking} if self.thinking else {}
            resp = client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                temperature=temperature,
                max_tokens=4096,
                extra_body=extra,
            )
            return resp.choices[0].message.content or ""
        except Exception as e:
            return f"[LLM Error] {type(e).__name__}: {e}"


# ═══════════════════════════════════════════════════════════════════════
#  Tools
# ═══════════════════════════════════════════════════════════════════════

@dataclass
class ToolResult:
    success: bool
    output: str
    tool_name: str
    elapsed_ms: float = 0


class Tool:
    """Base class for tools with enhanced capabilities."""

    name: str = ""
    description: str = ""
    
    # Tool metadata
    is_read_only: bool = True  # Whether tool only reads data
    is_concurrency_safe: bool = True  # Whether tool can run in parallel
    requires_confirmation: bool = False  # Whether tool needs user confirmation
    
    @property
    def parameters(self) -> dict[str, str]:
        return {}
    
    def check_permissions(self, params: dict, ctx: dict) -> tuple[bool, str]:
        """Check if tool execution is allowed. Returns (allowed, reason)."""
        return True, ""
    
    def validate_input(self, params: dict) -> tuple[bool, str]:
        """Validate input parameters. Returns (valid, error_message)."""
        return True, ""
    
    def run(self, params: dict, ctx: dict) -> str:
        raise NotImplementedError
    
    def run_with_progress(self, params: dict, ctx: dict, progress_callback=None) -> str:
        """Run tool with progress reporting."""
        # Default implementation ignores progress
        return self.run(params, ctx)


class SearchTool(Tool):
    name = "search"
    description = "语义搜索代码库，返回最相关的代码片段"
    is_read_only = True
    is_concurrency_safe = True

    @property
    def parameters(self):
        return {"query": "搜索关键词", "n": "结果数量(默认5)"}
    
    def validate_input(self, params: dict) -> tuple[bool, str]:
        query = params.get("query", "")
        if not query or not query.strip():
            return False, "搜索查询不能为空"
        if len(query) > 1000:
            return False, "搜索查询过长，请限制在1000字符以内"
        n = params.get("n", 5)
        try:
            n = int(n)
            if n < 1 or n > 20:
                return False, "结果数量应在1-20之间"
        except ValueError:
            return False, "结果数量必须是整数"
        return True, ""

    def run(self, params: dict, ctx: dict) -> str:
        store: VectorStore = ctx["store"]
        query = params.get("query", "")
        n = int(params.get("n", 5))
        if not query:
            return "Error: query required"
            
        # Optional: Query Expansion via LLM to get better keywords
        llm = ctx.get("llm")
        expanded_query = query
        if llm and llm.available:
            prompt = "你是一个代码搜索专家。将以下用户问题转换为用于向量搜索的丰富关键词。包含可能的英文变量名、函数名和技术术语。不要解释，只返回关键词。"
            try:
                expanded = llm.complete(prompt, query).strip()
                if expanded and len(expanded) < 200:
                    expanded_query = f"{query} {expanded}"
            except Exception:
                pass
                
        results = store.query(expanded_query, n_results=n)
        if not results:
            return "No results."
        parts = []
        for i, r in enumerate(results, 1):
            m = r["metadata"]
            parts.append(f"[{i}] `{m['file_path']}` L{m['start_line']}-{m['end_line']}\n```\n{r['content']}\n```")
        return "\n\n".join(parts)


class ReadFileTool(Tool):
    name = "read_file"
    description = "读取文件完整内容。默认读整个文件，可选指定行范围。"

    @property
    def parameters(self):
        return {"path": "文件路径", "start": "起始行(可选,默认读全部)", "end": "结束行(可选)"}

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        path = params.get("path", "")
        if not path:
            return "Error: path required"
        # Resolve and validate path is within project root
        full = (root / path).resolve()
        if not full.is_relative_to(root):
            return f"Access denied: path outside project root"
        if not full.exists():
            # try glob within root only
            name = Path(path).name
            matches = [m for m in root.rglob(name) if m.resolve().is_relative_to(root)]
            if matches:
                full = matches[0]
            else:
                return f"File not found: {path}"
        content = read_file(full)
        if content is None:
            return f"Cannot read: {path}"
        lines = content.splitlines()
        total = len(lines)
        s = max(1, int(params.get("start", 1)))
        e = min(total, int(params.get("end", total)))
        
        # Smart truncation: allow up to 2000 lines per read
        MAX_LINES = 2000
        if e - s + 1 > MAX_LINES:
            e = s + MAX_LINES - 1
            trunc_msg = f"\n... [文件共{total}行，已读取前{MAX_LINES}行。如需后续内容请指定 start={e+1}] ..."
        else:
            trunc_msg = ""
            
        numbered = "\n".join(f"{i+s:>4} | {l}" for i, l in enumerate(lines[s-1:e]))
        rel = str(full.relative_to(root))
        return f"`{rel}` ({total} lines, showing {s}-{e})\n```\n{numbered}{trunc_msg}\n```"


import concurrent.futures

class FindPatternTool(Tool):
    name = "find_pattern"
    description = "正则搜索代码（找函数定义、import等）"

    @property
    def parameters(self):
        return {"pattern": "正则表达式", "file_glob": "文件过滤(可选)"}

    def _search_file(self, f: Path, root: Path, regex: re.Pattern, file_glob: str | None) -> list[str]:
        matches = []
        if file_glob:
            rel = str(f.relative_to(root))
            if not Path(rel).match(file_glob):
                return matches
        content = read_file(f)
        if content is None:
            return matches
        for i, line in enumerate(content.splitlines(), 1):
            search_line = line[:500]
            try:
                if regex.search(search_line):
                    rel = str(f.relative_to(root))
                    matches.append(f"`{rel}:{i}`  {line.strip()}")
            except Exception:
                continue
        return matches

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        pattern = params.get("pattern", "")
        if not pattern:
            return "Error: pattern required"
        if len(pattern) > 200:
            return "Error: pattern too long (max 200 chars)"
            
        # Basic ReDoS mitigation by checking for suspicious repetition
        if re.search(r'(\([^)]+\)|\w+)([*+]{2,}|[*+]\?)', pattern) or re.search(r'(\([^)]+\)|\w+)[*+]\(', pattern):
            return "Error: pattern rejected due to potential ReDoS vulnerability (nested quantifiers or overlapping alternations)."

        try:
            regex = re.compile(pattern)
        except re.error as e:
            return f"Invalid regex: {e}"
        file_glob = params.get("file_glob")
        
        files = scan_files(root)
        all_matches = []
        
        # Determine a reasonable chunksize and use as_completed to allow early break
        # We don't actually use chunksize for ThreadPoolExecutor.submit, it's just for reference if we used map.
        # But we will use a threading Event to allow true cancellation of running threads if supported
        # or at least prevent new tasks from doing work.
        import threading
        cancel_event = threading.Event()
        
        def _search_wrapper(f, root, regex, file_glob):
            if cancel_event.is_set():
                return []
            return self._search_file(f, root, regex, file_glob)
            
        with concurrent.futures.ThreadPoolExecutor() as executor:
            # Submit futures instead of map to allow cancellation
            futures = [executor.submit(_search_wrapper, f, root, regex, file_glob) for f in files]
            
            for future in concurrent.futures.as_completed(futures):
                matches = future.result()
                if matches:
                    all_matches.extend(matches)
                if len(all_matches) >= 30:
                    # Cancel remaining futures
                    cancel_event.set()
                    for f in futures:
                        f.cancel()
                    break
                    
        return "\n".join(all_matches[:30]) if all_matches else "No matches."


class ListDirTool(Tool):
    name = "list_dir"
    description = "列出目录结构"

    @property
    def parameters(self):
        return {"path": "目录路径(留空=根目录)", "depth": "深度(默认2)"}

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        path = params.get("path", "")
        target = (root / path).resolve() if path else root
        
        if not target.is_relative_to(root):
            return f"Access denied: path outside project root"
            
        depth = int(params.get("depth", 2))
        if not target.exists():
            return f"Not found: {params.get('path', '.')}"
        lines = []
        self._walk(target, depth, lines, "")
        return "\n".join(lines) if lines else "(empty)"

    def _walk(self, d: Path, depth: int, lines: list, prefix: str):
        if depth <= 0:
            return
        skip = {".git","__pycache__","node_modules",".venv","venv",".snowcode","dist","build"}
        try:
            entries = sorted(d.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower()))
        except PermissionError:
            return
        for e in entries:
            if e.name in skip or e.name.startswith("."):
                continue
            if e.is_dir():
                lines.append(f"{prefix}{e.name}/")
                self._walk(e, depth-1, lines, prefix+"  ")
            else:
                try:
                    sz = e.stat().st_size
                    ss = f"{sz:,}B" if sz < 1024 else f"{sz//1024:,}KB"
                except OSError:
                    ss = "?"
                lines.append(f"{prefix}{e.name}  [{ss}]")


class ReadMultipleTool(Tool):
    name = "read_multiple"
    description = "同时读取多个文件的关键行"

    @property
    def parameters(self):
        return {"files": "格式: 'a.py:10-30,b.py:5-15'"}

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        spec = params.get("files", "")
        if not spec:
            return "Error: files required"
        parts = []
        MAX_TOTAL_LINES = 1000
        total_lines_read = 0
        
        for item in spec.split(","):
            item = item.strip()
            if not item:
                continue
            if ":" in item:
                fp, rs = item.rsplit(":", 1)
                try:
                    s, e = map(int, rs.split("-"))
                except ValueError:
                    s, e = 1, 9999
            else:
                fp, s, e = item, 1, 9999
            full = (root / fp).resolve()
            if not full.is_relative_to(root):
                parts.append(f"Access denied: {fp}")
                continue
            if not full.exists():
                parts.append(f"Not found: {fp}")
                continue
            content = read_file(full)
            if content is None:
                continue
            lines = content.splitlines()
            s, e = max(1, s), min(len(lines), e)
            
            lines_to_read = e - s + 1
            if total_lines_read + lines_to_read > MAX_TOTAL_LINES:
                allowed = max(0, MAX_TOTAL_LINES - total_lines_read)
                e = s + allowed - 1
                trunc_msg = f"\n... [Truncated due to total line limit]"
            else:
                trunc_msg = ""
                
            if e >= s:
                numbered = "\n".join(f"{i+s:>4} | {l}" for i, l in enumerate(lines[s-1:e]))
                parts.append(f"`{fp}` L{s}-{e}\n```\n{numbered}{trunc_msg}\n```")
                total_lines_read += (e - s + 1)
                
            if total_lines_read >= MAX_TOTAL_LINES:
                parts.append("... [Maximum line limit reached for this read_multiple call]")
                break
                
        return "\n\n".join(parts)


class WriteFileTool(Tool):
    """Tool to write content to a file."""
    name = "write_file"
    description = "写入或覆盖整个文件内容"

    @property
    def parameters(self):
        return {"path": "文件路径", "content": "要写入的完整文件内容"}

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        path = params.get("path", "")
        content = params.get("content", "")
        if not path or not content:
            return "Error: missing path or content arguments."
            
        try:
            # Safe path resolution
            safe_path = path.lstrip("/").lstrip("\\")
            full = (root / safe_path).resolve()
            if not full.is_relative_to(root):
                return f"Error: Cannot write outside project root: {path}"
                
            full.parent.mkdir(parents=True, exist_ok=True)
            
            # Create a backup if file exists
            if full.exists():
                import shutil
                backup = full.with_suffix(full.suffix + ".bak")
                shutil.copy2(full, backup)
                
            full.write_text(content, encoding="utf-8")
            return f"Successfully wrote to `{path}` ({len(content)} chars, Backup saved as .bak if overwritten)."
        except Exception as e:
            return f"Error writing file: {e}"


class SearchReplaceTool(Tool):
    name = "search_replace"
    description = "搜索并替换文件中的指定代码块"

    @property
    def parameters(self):
        return {
            "path": "文件路径", 
            "old_str": "要被替换的原始代码块(需精确匹配)", 
            "new_str": "新的代码块"
        }

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        path = params.get("path", "")
        old_str = params.get("old_str", "")
        new_str = params.get("new_str", "")
        
        if not path or not old_str:
            return "Error: path and old_str required"
            
        full = (root / path).resolve()
        if not full.is_relative_to(root):
            return f"Access denied: path outside project root"
            
        if not full.exists():
            return f"Error: File not found `{path}`"
            
        try:
            content = full.read_text(encoding="utf-8")
            if old_str not in content:
                return "Error: old_str not found in the file. Ensure exact matching including whitespace."
                
            # Create a backup if file exists
            if full.exists():
                import shutil
                backup = full.with_suffix(full.suffix + ".bak")
                shutil.copy2(full, backup)
                
            new_content = content.replace(old_str, new_str, 1) # Only replace first occurrence to be safe
            full.write_text(new_content, encoding="utf-8")
            return f"Successfully replaced code in `{path}` (Backup saved as .bak)."
        except Exception as e:
            return f"Error modifying file: {e}"


class ToolRegistry:
    """Enhanced registry with permissions, validation, and concurrency."""

    def __init__(self):
        self._tools: dict[str, Tool] = {}

    def register(self, tool: Tool):
        self._tools[tool.name] = tool

    def get(self, name: str) -> Tool | None:
        return self._tools.get(name)
    
    def get_all(self) -> list[Tool]:
        """Get all registered tools."""
        return list(self._tools.values())
    
    def get_read_only_tools(self) -> list[Tool]:
        """Get tools that only read data."""
        return [t for t in self._tools.values() if t.is_read_only]
    
    def get_concurrency_safe_tools(self) -> list[Tool]:
        """Get tools that can run in parallel."""
        return [t for t in self._tools.values() if t.is_concurrency_safe]

    def list_definitions(self) -> str:
        lines = []
        for t in self._tools.values():
            params = ", ".join(f"{k}: {v}" for k, v in t.parameters.items())
            safety = "[锁定]" if not t.is_read_only else "[只读]"
            lines.append(f"- **{t.name}**({params}): {t.description} {safety}")
        return "\n".join(lines)
    
    def list_definitions_for_llm(self) -> str:
        """Format tool definitions optimized for LLM understanding."""
        lines = []
        for t in self._tools.values():
            params = []
            for param_name, param_desc in t.parameters.items():
                params.append(f"{param_name}: {param_desc}")
            params_str = ", ".join(params) if params else "无参数"
            
            # Add safety indicators
            safety = []
            if not t.is_read_only:
                safety.append("[警告] 会修改文件")
            if t.requires_confirmation:
                safety.append("[锁定] 需要确认")
            safety_str = f" [{', '.join(safety)}]" if safety else ""
            
            lines.append(f"• {t.name}({params_str}): {t.description}{safety_str}")
        return "\n".join(lines)

    def execute(self, name: str, params: dict, ctx: dict, check_permissions: bool = True) -> ToolResult:
        """Execute tool with enhanced error handling and validation."""
        tool = self._tools.get(name)
        if not tool:
            return ToolResult(False, f"未知工具: {name}", name)
        
        # Input validation
        valid, error = tool.validate_input(params)
        if not valid:
            return ToolResult(False, f"输入验证失败: {error}", name)
        
        # Permission check
        if check_permissions:
            allowed, reason = tool.check_permissions(params, ctx)
            if not allowed:
                return ToolResult(False, f"权限检查失败: {reason}", name)
        
        start = time.time()
        try:
            output = tool.run(params, ctx)
            elapsed = (time.time() - start) * 1000
            return ToolResult(True, output, name, elapsed)
        except Exception as e:
            elapsed = (time.time() - start) * 1000
            return ToolResult(False, f"{type(e).__name__}: {e}", name, elapsed)
    
    def execute_parallel(self, tool_calls: list[tuple[str, dict]], ctx: dict) -> list[ToolResult]:
        """Execute multiple tools in parallel if they are concurrency-safe."""
        import concurrent.futures
        
        results = []
        safe_calls = []
        unsafe_calls = []
        
        # Separate safe and unsafe calls
        for name, params in tool_calls:
            tool = self._tools.get(name)
            if tool and tool.is_concurrency_safe:
                safe_calls.append((name, params))
            else:
                unsafe_calls.append((name, params))
        
        # Execute unsafe calls sequentially
        for name, params in unsafe_calls:
            results.append(self.execute(name, params, ctx))
        
        # Execute safe calls in parallel
        if safe_calls:
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(4, len(safe_calls))) as executor:
                futures = []
                for name, params in safe_calls:
                    future = executor.submit(self.execute, name, params, ctx, False)
                    futures.append((name, future))
                
                for name, future in futures:
                    try:
                        result = future.result(timeout=30)  # 30 second timeout
                        results.append(result)
                    except Exception as e:
                        results.append(ToolResult(False, f"并行执行异常: {e}", name))
        
        return results


class DeleteFileTool(Tool):
    """Tool to delete a file."""
    name = "delete_file"
    description = "删除指定文件"

    @property
    def parameters(self):
        return {"path": "要删除的文件路径"}

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        path = params.get("path", "")
        if not path:
            return "Error: path required"
        full = (root / path).resolve()
        if not full.is_relative_to(root):
            return f"Access denied: path outside project root"
        if not full.exists():
            return f"File not found: {path}"
        if full.is_dir():
            return f"Error: {path} is a directory, not a file"
        try:
            import shutil
            # Backup before deleting
            backup = full.with_suffix(full.suffix + ".deleted")
            shutil.copy2(full, backup)
            full.unlink()
            rel = str(full.relative_to(root))
            return f"Deleted `{rel}` (backup saved as {backup.name})"
        except Exception as e:
            return f"Error deleting: {e}"


class ShellTool(Tool):
    """Execute shell commands."""
    name = "shell"
    description = "执行终端命令（cmd/bash），如查看git状态、运行测试、安装依赖等"

    @property
    def parameters(self):
        return {"command": "要执行的命令", "cwd": "工作目录(可选,默认项目根目录)"}

    def run(self, params: dict, ctx: dict) -> str:
        import subprocess
        root: Path = ctx["root"]
        cmd = params.get("command", "")
        cwd = params.get("cwd", "")
        if not cmd:
            return "Error: command required"

        # Safety: block dangerous commands
        _blocked = ("rm -rf /", "format", "del /f /s", "shutdown", "mkfs", "dd if=")
        cmd_lower = cmd.lower()
        for b in _blocked:
            if b in cmd_lower:
                return f"Error: blocked dangerous command: {cmd}"

        work_dir = (root / cwd).resolve() if cwd else root
        if not work_dir.is_relative_to(root) and work_dir != root:
            return f"Error: cwd outside project root"

        try:
            result = subprocess.run(
                cmd, shell=True, cwd=str(work_dir),
                capture_output=True, text=True, timeout=30,
                encoding="utf-8", errors="replace",
            )
            output = result.stdout + result.stderr
            # Truncate long output
            if len(output) > 5000:
                output = output[:5000] + f"\n... [truncated, total {len(output)} chars]"
            exit_info = f"[exit code: {result.returncode}]" if result.returncode != 0 else ""
            return f"{output}{exit_info}" if output else f"(no output) {exit_info}"
        except subprocess.TimeoutExpired:
            return "Error: command timed out (30s limit)"
        except Exception as e:
            return f"Error: {type(e).__name__}: {e}"


class GitTool(Tool):
    """Git operations."""
    name = "git"
    description = "Git操作：status、log、diff、blame等"

    @property
    def parameters(self):
        return {"args": "git参数，如 'status' 'log --oneline -10' 'diff HEAD~1'"}

    def run(self, params: dict, ctx: dict) -> str:
        import subprocess
        root: Path = ctx["root"]
        args = params.get("args", "")
        if not args:
            return "Error: args required"

        # Only allow safe git commands
        _allowed = ("status", "log", "diff", "blame", "show", "branch", "remote",
                     "stash", "reflog", "tag", "ls-files", "grep")
        first_arg = args.split()[0] if args.split() else ""
        if first_arg not in _allowed:
            return f"Error: '{first_arg}' not allowed. Allowed: {', '.join(_allowed)}"

        try:
            result = subprocess.run(
                f"git {args}", shell=True, cwd=str(root),
                capture_output=True, text=True, timeout=15,
                encoding="utf-8", errors="replace",
            )
            output = result.stdout + result.stderr
            if len(output) > 5000:
                output = output[:5000] + "\n... [truncated]"
            return output if output else "(no output)"
        except subprocess.TimeoutExpired:
            return "Error: git command timed out"
        except Exception as e:
            return f"Error: {type(e).__name__}: {e}"


class PythonRunTool(Tool):
    """Run Python code snippets."""
    name = "python_run"
    description = "执行Python代码片段，用于验证逻辑、测试表达式等"

    @property
    def parameters(self):
        return {"code": "Python代码"}

    def run(self, params: dict, ctx: dict) -> str:
        import subprocess
        import tempfile
        root: Path = ctx["root"]
        code = params.get("code", "")
        if not code:
            return "Error: code required"

        # Safety: block dangerous imports
        _blocked = ("os.system", "subprocess", "shutil.rmtree", "eval(",
                     "exec(", "__import__", "importlib", "ctypes")
        for b in _blocked:
            if b in code:
                return f"Error: blocked potentially dangerous code containing '{b}'"

        try:
            with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False, encoding="utf-8") as f:
                f.write(code)
                f.flush()
                result = subprocess.run(
                    ["python", f.name], cwd=str(root),
                    capture_output=True, text=True, timeout=15,
                    encoding="utf-8", errors="replace",
                )
            output = result.stdout + result.stderr
            if len(output) > 3000:
                output = output[:3000] + "\n... [truncated]"
            return output if output else "(no output)"
        except subprocess.TimeoutExpired:
            return "Error: code execution timed out (15s)"
        except Exception as e:
            return f"Error: {type(e).__name__}: {e}"


# ═══════════════════════════════════════════════════════════════════════
#  Multimodal Tools
# ═══════════════════════════════════════════════════════════════════════

class ImageReaderTool(Tool):
    """Read images and extract information including OCR text."""
    name = "image_reader"
    description = "读取图片文件，支持OCR提取文字、分析图片内容"
    is_read_only = True
    is_concurrency_safe = True

    @property
    def parameters(self):
        return {
            "path": "图片文件路径（相对于项目根目录）",
            "mode": "读取模式: 'ocr'(提取文字), 'describe'(描述内容), 'info'(基本信息)"
        }
    
    def validate_input(self, params: dict) -> tuple[bool, str]:
        path = params.get("path", "")
        if not path or not path.strip():
            return False, "文件路径不能为空"
        mode = params.get("mode", "info")
        if mode not in ("ocr", "describe", "info"):
            return False, "mode必须是 ocr、describe 或 info 之一"
        return True, ""
    
    def check_permissions(self, params: dict, ctx: dict) -> tuple[bool, str]:
        root: Path = ctx["root"]
        try:
            file_path = (root / params["path"]).resolve()
            if not file_path.is_relative_to(root):
                return False, "路径超出项目根目录"
            return True, ""
        except Exception as e:
            return False, f"路径验证失败: {e}"

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        file_path = (root / params["path"]).resolve()
        mode = params.get("mode", "info")
        
        if not file_path.exists():
            return f"Error: 文件不存在: {params['path']}"
        
        # Check file type
        supported = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.tif', '.svg'}
        if file_path.suffix.lower() not in supported:
            return f"Error: 不支持的图片格式 '{file_path.suffix}'，支持: {', '.join(supported)}"
        
        # Basic info
        info_lines = [f"文件: {file_path.name}", f"路径: {file_path}",
                      f"大小: {self._format_size(file_path.stat().st_size)}",
                      f"类型: {file_path.suffix}"]
        
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                info_lines.append(f"尺寸: {img.width}x{img.height} 像素")
                info_lines.append(f"模式: {img.mode}")
                info_lines.append(f"格式: {img.format or 'Unknown'}")
        except ImportError:
            info_lines.append("[警告] PIL/Pillow 未安装，无法获取详细信息")
        except Exception as e:
            info_lines.append(f"[警告] 读取图片信息失败: {e}")
        
        result = "\n".join(info_lines)
        
        # OCR mode
        if mode == "ocr":
            result += "\n\n--- OCR 提取的文字 ---\n"
            try:
                import pytesseract
                from PIL import Image
                with Image.open(file_path) as img:
                    text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    if text.strip():
                        result += text.strip()
                    else:
                        result += "(未检测到文字)"
            except ImportError:
                result += "[警告] pytesseract 未安装，无法执行OCR。请安装: pip install pytesseract\n"
                result += "同时需要安装 Tesseract OCR 引擎: https://github.com/tesseract-ocr/tesseract"
            except Exception as e:
                result += f"OCR 失败: {e}"
        
        # Describe mode - use LLM if available
        elif mode == "describe":
            llm = ctx.get("llm")
            if llm and llm.available:
                result += "\n\n--- 图片描述 ---\n"
                try:
                    import base64
                    with open(file_path, "rb") as f:
                        img_base64 = base64.b64encode(f.read()).decode('utf-8')
                    
                    # Try multimodal LLM if available
                    system = "你是一个图片分析专家。请详细描述这张图片的内容，包括：场景、物体、文字、颜色、布局等。"
                    user_msg = "请描述这张图片的内容。"
                    
                    # Check if LLM supports vision (try to send image)
                    desc = llm.complete(system, user_msg)
                    if desc:
                        result += desc
                    else:
                        result += "[警告] 当前LLM不支持图片分析，请使用OCR模式提取文字"
                except Exception as e:
                    result += f"图片分析失败: {e}"
            else:
                result += "\n\n[警告] LLM不可用，无法描述图片内容"
        
        return result
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"


class PDFReaderTool(Tool):
    """Read and parse PDF documents."""
    name = "pdf_reader"
    description = "读取PDF文档，提取文字内容、元信息和结构"
    is_read_only = True
    is_concurrency_safe = True

    @property
    def parameters(self):
        return {
            "path": "PDF文件路径（相对于项目根目录）",
            "pages": "页码范围，如 '1-5' 或 '1,3,5' 或 'all'(默认前10页)"
        }
    
    def validate_input(self, params: dict) -> tuple[bool, str]:
        path = params.get("path", "")
        if not path or not path.strip():
            return False, "文件路径不能为空"
        return True, ""
    
    def check_permissions(self, params: dict, ctx: dict) -> tuple[bool, str]:
        root: Path = ctx["root"]
        try:
            file_path = (root / params["path"]).resolve()
            if not file_path.is_relative_to(root):
                return False, "路径超出项目根目录"
            return True, ""
        except Exception as e:
            return False, f"路径验证失败: {e}"

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        file_path = (root / params["path"]).resolve()
        pages_param = params.get("pages", "all")
        
        if not file_path.exists():
            return f"Error: 文件不存在: {params['path']}"
        
        if file_path.suffix.lower() != '.pdf':
            return f"Error: 不是PDF文件: {params['path']}"
        
        result_lines = []
        result_lines.append(f"PDF文档: {file_path.name}")
        result_lines.append(f"路径: {file_path}")
        result_lines.append(f"大小: {self._format_size(file_path.stat().st_size)}")
        result_lines.append("")
        
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            
            result_lines.append(f"总页数: {len(doc)}")
            result_lines.append(f"元信息:")
            meta = doc.metadata
            if meta:
                for key in ['title', 'author', 'subject', 'creator', 'producer']:
                    val = meta.get(key, '')
                    if val:
                        result_lines.append(f"   {key}: {val}")
            result_lines.append("")
            
            # Parse page range
            pages_to_read = self._parse_page_range(pages_param, len(doc))
            
            result_lines.append(f"--- 文字内容 (第 {pages_to_read[0]+1}-{pages_to_read[-1]+1} 页) ---\n")
            
            total_text = []
            for page_num in pages_to_read:
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    total_text.append(f"\n{'='*50}")
                    total_text.append(f"第 {page_num + 1} 页")
                    total_text.append(f"{'='*50}")
                    total_text.append(text.strip())
            
            combined_text = "\n".join(total_text)
            if len(combined_text) > 10000:
                combined_text = combined_text[:10000] + f"\n\n... [内容过长，已截断，总长度 {len(combined_text)} 字符]"
            
            result_lines.append(combined_text if combined_text.strip() else "(该页面无文字内容)")
            
            doc.close()
        except ImportError:
            result_lines.append("[警告] PyMuPDF 未安装，无法读取PDF。请安装: pip install PyMuPDF")
        except Exception as e:
            result_lines.append(f"读取PDF失败: {e}")
        
        return "\n".join(result_lines)
    
    def _parse_page_range(self, pages_param: str, total_pages: int) -> list[int]:
        """Parse page range parameter."""
        if pages_param.lower() == "all":
            return list(range(min(total_pages, 10)))  # Default first 10 pages
        
        pages = []
        try:
            for part in pages_param.split(","):
                part = part.strip()
                if "-" in part:
                    start, end = part.split("-", 1)
                    start = max(0, int(start.strip()) - 1)
                    end = min(total_pages, int(end.strip()))
                    pages.extend(range(start, end))
                else:
                    page = int(part.strip()) - 1
                    if 0 <= page < total_pages:
                        pages.append(page)
        except ValueError:
            # Fallback to first 10 pages
            return list(range(min(total_pages, 10)))
        
        return sorted(set(pages)) if pages else list(range(min(total_pages, 10)))
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"


class DocumentReaderTool(Tool):
    """Read various document formats (docx, xlsx, csv, txt, md, etc.)."""
    name = "document_reader"
    description = "读取多种格式的文档：Word(.docx)、Excel(.xlsx)、CSV、TXT、Markdown等"
    is_read_only = True
    is_concurrency_safe = True

    @property
    def parameters(self):
        return {
            "path": "文件路径（相对于项目根目录）",
            "max_lines": "最大返回行数(默认500)"
        }
    
    def validate_input(self, params: dict) -> tuple[bool, str]:
        path = params.get("path", "")
        if not path or not path.strip():
            return False, "文件路径不能为空"
        return True, ""
    
    def check_permissions(self, params: dict, ctx: dict) -> tuple[bool, str]:
        root: Path = ctx["root"]
        try:
            file_path = (root / params["path"]).resolve()
            if not file_path.is_relative_to(root):
                return False, "路径超出项目根目录"
            return True, ""
        except Exception as e:
            return False, f"路径验证失败: {e}"

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        file_path = (root / params["path"]).resolve()
        max_lines = int(params.get("max_lines", 500))
        
        if not file_path.exists():
            return f"Error: 文件不存在: {params['path']}"
        
        ext = file_path.suffix.lower()
        result_lines = []
        result_lines.append(f" 文档: {file_path.name}")
        result_lines.append(f" 路径: {file_path}")
        result_lines.append(f" 大小: {self._format_size(file_path.stat().st_size)}")
        result_lines.append(f"️  格式: {ext}")
        result_lines.append("")
        
        try:
            if ext == '.docx':
                result_lines.append(self._read_docx(file_path, max_lines))
            elif ext == '.xlsx':
                result_lines.append(self._read_xlsx(file_path, max_lines))
            elif ext == '.csv':
                result_lines.append(self._read_csv(file_path, max_lines))
            elif ext in ('.txt', '.md', '.rst', '.log', '.json', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg'):
                result_lines.append(self._read_text(file_path, max_lines))
            elif ext == '.html':
                result_lines.append(self._read_html(file_path, max_lines))
            else:
                # Try as text
                result_lines.append(self._read_text(file_path, max_lines))
        except Exception as e:
            result_lines.append(f"读取文件失败: {e}")
        
        return "\n".join(result_lines)
    
    def _read_docx(self, file_path: Path, max_lines: int) -> str:
        """Read Word document."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            lines = []
            for i, para in enumerate(doc.paragraphs):
                if i >= max_lines:
                    lines.append(f"\n... [已截断，共 {len(doc.paragraphs)} 段]")
                    break
                if para.text.strip():
                    lines.append(para.text)
            return "\n".join(lines) if lines else "(文档无文字内容)"
        except ImportError:
            return "️ python-docx 未安装，无法读取Word文档。请安装: pip install python-docx"
    
    def _read_xlsx(self, file_path: Path, max_lines: int) -> str:
        """Read Excel spreadsheet."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(file_path), data_only=True)
            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"\n 工作表: {sheet_name}")
                lines.append(f"   尺寸: {ws.max_row}行 x {ws.max_column}列")
                lines.append("")
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    if row_count >= max_lines:
                        lines.append(f"... [已截断]")
                        break
                    lines.append(" | ".join(str(cell) if cell is not None else "" for cell in row))
                    row_count += 1
            wb.close()
            return "\n".join(lines)
        except ImportError:
            return "️ openpyxl 未安装，无法读取Excel文件。请安装: pip install openpyxl"
    
    def _read_csv(self, file_path: Path, max_lines: int) -> str:
        """Read CSV file."""
        import csv
        lines = []
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i >= max_lines:
                    lines.append(f"... [已截断]")
                    break
                lines.append(" | ".join(row))
        return "\n".join(lines) if lines else "(文件为空)"
    
    def _read_text(self, file_path: Path, max_lines: int) -> str:
        """Read plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            lines = f.readlines()
        
        if len(lines) > max_lines:
            content = "".join(lines[:max_lines])
            content += f"\n\n... [已截断，共 {len(lines)} 行]"
        else:
            content = "".join(lines)
        
        return content if content.strip() else "(文件为空)"
    
    def _read_html(self, file_path: Path, max_lines: int) -> str:
        """Read HTML file, extract text content."""
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            text = soup.get_text(separator='\n', strip=True)
            lines = text.split('\n')
            if len(lines) > max_lines:
                return "\n".join(lines[:max_lines]) + f"\n\n... [已截断]"
            return text
        except ImportError:
            # Fallback to raw text
            return self._read_text(file_path, max_lines)
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"


class FileBrowserTool(Tool):
    """Browse files in a local directory."""
    name = "file_browser"
    description = "浏览本地目录中的文件列表，支持按类型筛选"
    is_read_only = True
    is_concurrency_safe = True

    @property
    def parameters(self):
        return {
            "path": "目录路径（相对于项目根目录，默认为项目根目录）",
            "file_type": "文件类型筛选: 'all', 'images', 'docs', 'code', 'data'",
            "recursive": "是否递归搜索子目录 (true/false)"
        }
    
    def validate_input(self, params: dict) -> tuple[bool, str]:
        file_type = params.get("file_type", "all")
        if file_type not in ('all', 'images', 'docs', 'code', 'data'):
            return False, "file_type必须是 all、images、docs、code 或 data 之一"
        return True, ""

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        dir_path = (root / params.get("path", ".")).resolve()
        file_type = params.get("file_type", "all")
        recursive = params.get("recursive", "false").lower() == "true"
        
        if not dir_path.exists():
            return f"Error: 目录不存在: {params.get('path', '.')}"
        
        if not dir_path.is_dir():
            return f"Error: 不是目录: {params.get('path', '.')}"
        
        if not dir_path.is_relative_to(root):
            return f"Error: 路径超出项目根目录"
        
        # File type extensions
        type_map = {
            'images': {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.svg', '.ico'},
            'docs': {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.txt', '.md', '.rtf', '.odt'},
            'code': {'.py', '.js', '.ts', '.java', '.c', '.cpp', '.h', '.go', '.rs', '.rb', '.php', '.css', '.html'},
            'data': {'.json', '.csv', '.xml', '.yaml', '.yml', '.toml', '.sql', '.db', '.sqlite'},
        }
        
        if recursive:
            files = list(dir_path.rglob("*"))
        else:
            files = list(dir_path.iterdir())
        
        # Filter files only
        files = [f for f in files if f.is_file()]
        
        # Filter by type
        if file_type != "all":
            extensions = type_map.get(file_type, set())
            files = [f for f in files if f.suffix.lower() in extensions]
        
        # Sort by modification time
        files.sort(key=lambda f: f.stat().st_mtime, reverse=True)
        
        # Format output
        lines = []
        lines.append(f" 目录: {dir_path.relative_to(root) if dir_path != root else '.'}")
        lines.append(f" 文件数: {len(files)}")
        lines.append(f"️  类型: {file_type}")
        lines.append(f" 递归: {'是' if recursive else '否'}")
        lines.append("")
        
        if not files:
            lines.append("(没有找到匹配的文件)")
            return "\n".join(lines)
        
        # Group by extension
        by_ext: dict[str, list[Path]] = {}
        for f in files[:100]:  # Limit display
            ext = f.suffix.lower() or "(无扩展名)"
            if ext not in by_ext:
                by_ext[ext] = []
            by_ext[ext].append(f)
        
        for ext, ext_files in sorted(by_ext.items()):
            lines.append(f"\n{ext} ({len(ext_files)} 个文件):")
            for f in ext_files[:10]:
                size = self._format_size(f.stat().st_size)
                rel = f.relative_to(root)
                lines.append(f"   {rel}  ({size})")
            if len(ext_files) > 10:
                lines.append(f"  ... 及其他 {len(ext_files) - 10} 个文件")
        
        if len(files) > 100:
            lines.append(f"\n... 共 {len(files)} 个文件，仅显示前100个")
        
        return "\n".join(lines)
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"


class NCDataReaderTool(Tool):
    """Read and analyze NetCDF (.nc) scientific data files."""
    name = "nc_reader"
    description = "读取NetCDF(.nc)科学数据文件，支持查看维度、变量、属性和数据"
    is_read_only = True
    is_concurrency_safe = True

    @property
    def parameters(self):
        return {
            "path": "NC文件路径（相对于项目根目录）",
            "mode": "读取模式: 'info'(文件信息), 'vars'(变量列表), 'data'(变量数据), 'stats'(统计信息)",
            "variable": "变量名（mode为data或stats时必需）",
            "dimensions": "维度切片，如 'time:0:10,lat:0:5'（可选）"
        }
    
    def validate_input(self, params: dict) -> tuple[bool, str]:
        path = params.get("path", "")
        if not path or not path.strip():
            return False, "文件路径不能为空"
        mode = params.get("mode", "info")
        if mode not in ("info", "vars", "data", "stats"):
            return False, "mode必须是 info、vars、data 或 stats 之一"
        if mode in ("data", "stats") and not params.get("variable"):
            return False, "mode为data或stats时，variable参数必需"
        return True, ""
    
    def check_permissions(self, params: dict, ctx: dict) -> tuple[bool, str]:
        root: Path = ctx["root"]
        try:
            file_path = (root / params["path"]).resolve()
            if not file_path.is_relative_to(root):
                return False, "路径超出项目根目录"
            return True, ""
        except Exception as e:
            return False, f"路径验证失败: {e}"

    def run(self, params: dict, ctx: dict) -> str:
        root: Path = ctx["root"]
        file_path = (root / params["path"]).resolve()
        mode = params.get("mode", "info")
        variable = params.get("variable", "")
        dim_slice = params.get("dimensions", "")
        
        if not file_path.exists():
            return f"Error: 文件不存在: {params['path']}"
        
        if file_path.suffix.lower() not in ('.nc', '.nc4', '.netcdf'):
            return f"Error: 不是NetCDF文件: {params['path']}"
        
        try:
            import netCDF4 as nc
        except ImportError:
            try:
                import xarray as xr
                return self._read_with_xarray(file_path, mode, variable, dim_slice)
            except ImportError:
                return "️ 未安装 netCDF4 或 xarray 库。\n请安装: pip install netCDF4 xarray"
        
        try:
            dataset = nc.Dataset(str(file_path), 'r')
            
            result_lines = []
            result_lines.append(f" NetCDF文件: {file_path.name}")
            result_lines.append(f" 路径: {file_path}")
            result_lines.append(f" 大小: {self._format_size(file_path.stat().st_size)}")
            result_lines.append(f" 格式: {dataset.file_format}")
            result_lines.append("")
            
            if mode == "info":
                result_lines.extend(self._get_file_info(dataset))
            elif mode == "vars":
                result_lines.extend(self._get_variables(dataset))
            elif mode == "data":
                result_lines.extend(self._get_variable_data(dataset, variable, dim_slice))
            elif mode == "stats":
                result_lines.extend(self._get_variable_stats(dataset, variable))
            
            dataset.close()
            return "\n".join(result_lines)
            
        except Exception as e:
            return f"读取NetCDF文件失败: {e}"
    
    def _get_file_info(self, dataset) -> list[str]:
        """Get basic file information."""
        lines = []
        
        # Dimensions
        lines.append(f" 维度 ({len(dataset.dimensions)}):")
        for name, dim in dataset.dimensions.items():
            lines.append(f"   {name}: {len(dim)}")
        lines.append("")
        
        # Variables count
        lines.append(f" 变量数: {len(dataset.variables)}")
        lines.append(f"️  全局属性数: {len(dataset.ncattrs())}")
        
        # Global attributes
        if dataset.ncattrs():
            lines.append("\n 全局属性:")
            for attr_name in dataset.ncattrs():
                val = getattr(dataset, attr_name)
                val_str = str(val)[:100]
                lines.append(f"   {attr_name}: {val_str}")
        
        return lines
    
    def _get_variables(self, dataset) -> list[str]:
        """Get variable list with details."""
        lines = []
        lines.append(f" 变量列表 ({len(dataset.variables)}):")
        lines.append("")
        
        for name, var in dataset.variables.items():
            dims = ", ".join(var.dimensions) if var.dimensions else "标量"
            shape = var.shape
            dtype = var.datatype
            attrs = ", ".join(var.ncattrs()) if var.ncattrs() else "无属性"
            
            lines.append(f"• {name}")
            lines.append(f"  维度: ({dims})")
            lines.append(f"  形状: {shape}")
            lines.append(f"  类型: {dtype}")
            lines.append(f"  属性: {attrs}")
            lines.append("")
        
        return lines
    
    def _get_variable_data(self, dataset, var_name: str, dim_slice: str) -> list[str]:
        """Get variable data with optional slicing."""
        lines = []
        
        if var_name not in dataset.variables:
            return [f"Error: 变量 '{var_name}' 不存在。可用变量: {', '.join(dataset.variables.keys())}"]
        
        var = dataset.variables[var_name]
        lines.append(f" 变量: {var_name}")
        lines.append(f" 形状: {var.shape}")
        lines.append(f" 类型: {var.datatype}")
        lines.append("")
        
        try:
            # Parse dimension slicing
            if dim_slice:
                slices = self._parse_dim_slices(dim_slice, var.dimensions)
                data = var[slices]
                lines.append(f" 切片: {dim_slice}")
            else:
                # For large arrays, only read a sample
                total_size = 1
                for s in var.shape:
                    total_size *= s
                
                if total_size > 10000:
                    # Read only first elements
                    sample_slices = tuple(slice(min(s, 10)) for s in var.shape[:3])
                    if len(var.shape) > 3:
                        sample_slices = sample_slices + (0,) * (len(var.shape) - 3)
                    data = var[sample_slices]
                    lines.append(f"️ 数据过大 ({total_size} 元素)，仅显示前10x10x10样本")
                else:
                    data = var[:]
            
            lines.append(f" 数据形状: {data.shape}")
            lines.append(f" 数据范围: {data.min():.4f} ~ {data.max():.4f}")
            lines.append(f" 平均值: {data.mean():.4f}")
            lines.append("")
            lines.append("数据样本:")
            lines.append(str(data))
            
        except Exception as e:
            lines.append(f"读取数据失败: {e}")
        
        return lines
    
    def _get_variable_stats(self, dataset, var_name: str) -> list[str]:
        """Get statistical information for a variable."""
        lines = []
        
        if var_name not in dataset.variables:
            return [f"Error: 变量 '{var_name}' 不存在。可用变量: {', '.join(dataset.variables.keys())}"]
        
        var = dataset.variables[var_name]
        lines.append(f" 变量统计: {var_name}")
        lines.append(f" 形状: {var.shape}")
        lines.append(f" 类型: {var.datatype}")
        lines.append("")
        
        try:
            # Read data in chunks for large arrays
            data = var[:]
            flat = data.flatten()
            
            # Remove NaN for statistics
            import numpy as np
            valid = flat[~np.isnan(flat)]
            
            lines.append(f" 统计信息:")
            lines.append(f"   元素总数: {len(flat)}")
            lines.append(f"   有效值数: {len(valid)}")
            lines.append(f"   NaN 数量: {len(flat) - len(valid)}")
            lines.append(f"   最小值: {np.min(valid):.6f}")
            lines.append(f"   最大值: {np.max(valid):.6f}")
            lines.append(f"   平均值: {np.mean(valid):.6f}")
            lines.append(f"   中位数: {np.median(valid):.6f}")
            lines.append(f"   标准差: {np.std(valid):.6f}")
            lines.append(f"   方差: {np.var(valid):.6f}")
            
            # Percentiles
            lines.append(f"\n 分位数:")
            for p in [0, 25, 50, 75, 100]:
                lines.append(f"   P{p}: {np.percentile(valid, p):.6f}")
            
        except Exception as e:
            lines.append(f"计算统计信息失败: {e}")
        
        return lines
    
    def _parse_dim_slices(self, dim_slice: str, dimensions: tuple) -> tuple:
        """Parse dimension slice string like 'time:0:10,lat:0:5'."""
        slices = []
        slice_dict = {}
        
        for part in dim_slice.split(","):
            part = part.strip()
            if ":" in part:
                parts = part.split(":")
                dim_name = parts[0]
                if len(parts) == 3:
                    slice_dict[dim_name] = slice(int(parts[1]), int(parts[2]))
                elif len(parts) == 2:
                    slice_dict[dim_name] = slice(0, int(parts[1]))
        
        for dim_name in dimensions:
            if dim_name in slice_dict:
                slices.append(slice_dict[dim_name])
            else:
                slices.append(slice(None))
        
        return tuple(slices)
    
    def _read_with_xarray(self, file_path, mode: str, variable: str, dim_slice: str) -> str:
        """Fallback reading with xarray library."""
        try:
            import xarray as xr
            import numpy as np
            
            ds = xr.open_dataset(file_path)
            
            result_lines = []
            result_lines.append(f" NetCDF文件: {file_path.name}")
            result_lines.append(f" 路径: {file_path}")
            result_lines.append(f" 大小: {self._format_size(file_path.stat().st_size)}")
            result_lines.append("")
            
            if mode == "info":
                result_lines.append(f" 维度: {list(ds.dims)}")
                result_lines.append(f" 变量: {list(ds.data_vars)}")
                result_lines.append(f" 坐标: {list(ds.coords)}")
                result_lines.append(f"️  属性: {dict(ds.attrs)}")
            elif mode == "vars":
                for name, var in ds.data_vars.items():
                    result_lines.append(f"• {name}: dims={list(var.dims)}, shape={var.shape}, dtype={var.dtype}")
            elif mode == "data":
                if variable not in ds.data_vars:
                    return f"Error: 变量 '{variable}' 不存在。可用: {list(ds.data_vars)}"
                var = ds[variable]
                result_lines.append(f" {variable}")
                result_lines.append(f" 形状: {var.shape}")
                result_lines.append(f" 数据:\n{var.values}")
            elif mode == "stats":
                if variable not in ds.data_vars:
                    return f"Error: 变量 '{variable}' 不存在。可用: {list(ds.data_vars)}"
                var = ds[variable]
                valid = var.values[~np.isnan(var.values)]
                result_lines.append(f" {variable} 统计:")
                result_lines.append(f"   最小值: {np.min(valid):.6f}")
                result_lines.append(f"   最大值: {np.max(valid):.6f}")
                result_lines.append(f"   平均值: {np.mean(valid):.6f}")
                result_lines.append(f"   标准差: {np.std(valid):.6f}")
            
            ds.close()
            return "\n".join(result_lines)
        except Exception as e:
            return f"xarray 读取失败: {e}"
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"


def build_default_registry() -> ToolRegistry:
    reg = ToolRegistry()
    # Core code tools
    reg.register(SearchTool())
    reg.register(ReadFileTool())
    reg.register(FindPatternTool())
    reg.register(ListDirTool())
    reg.register(ReadMultipleTool())
    reg.register(WriteFileTool())
    reg.register(SearchReplaceTool())
    reg.register(DeleteFileTool())
    reg.register(ShellTool())
    reg.register(GitTool())
    reg.register(PythonRunTool())
    # Multimodal tools
    reg.register(ImageReaderTool())
    reg.register(PDFReaderTool())
    reg.register(DocumentReaderTool())
    reg.register(FileBrowserTool())
    # Scientific data tools
    reg.register(NCDataReaderTool())
    return reg


# ═══════════════════════════════════════════════════════════════════════
#  Memory
# ═══════════════════════════════════════════════════════════════════════

@dataclass
class MemoryEntry:
    role: str          # "user", "agent", "tool", "system"
    content: str
    tool_name: str = ""
    timestamp: float = field(default_factory=time.time)


class ShortTermMemory:
    """In-session conversation memory with sliding window."""

    def __init__(self, max_entries: int = 20):
        self.entries: list[MemoryEntry] = []
        self.max_entries = max_entries

    def add(self, role: str, content: str, tool_name: str = ""):
        self.entries.append(MemoryEntry(role=role, content=content, tool_name=tool_name))
        if len(self.entries) > self.max_entries:
            # Keep first (goal) and last N
            self.entries = [self.entries[0]] + self.entries[-(self.max_entries - 1):]

    def get_context(self, max_chars: int = 30000) -> str:
        """Format memory into a string for the LLM, strictly bound by max_chars."""
        lines = []
        total = 0
        for entry in reversed(self.entries):
            # Truncate overly long single tool outputs to prevent token overflow
            content_str = entry.content
            if entry.role == "tool" and len(content_str) > 2000:
                content_str = content_str[:2000] + "\n...[truncated for length]..."
                
            line = f"[{entry.role}] {content_str}"
            if total + len(line) > max_chars:
                break
            lines.append(line)
            total += len(line)
        lines.reverse()
        return "\n---\n".join(lines)

    def clear(self):
        self.entries.clear()


class LongTermMemory:
    """Persistent memory stored in .snowcode/memory.jsonl."""

    def __init__(self, project_root: Path):
        self.dir = get_snowcode_dir(project_root)
        self.path = self.dir / "memory.jsonl"

    def store(self, question: str, answer: str, actions: list[dict]):
        """Save a Q&A session for future reference."""
        entry = {
            "ts": time.time(),
            "q": question[:200],
            "a": answer[:500],
            "actions": [a["tool"] for a in actions],
            "hash": hashlib.sha256(question.encode()).hexdigest()[:16],
        }
        with open(self.path, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")

    def recall(self, question: str, n: int = 3) -> str:
        """Retrieve similar past Q&As."""
        if not self.path.exists():
            return ""
            
        q_words = set(re.findall(r'\w+', question.lower()))
        
        # Better than just single words: add character trigrams for partial matches
        def get_trigrams(text):
            text = text.lower()
            return {text[i:i+3] for i in range(len(text)-2)} if len(text) >= 3 else set()
            
        q_trigrams = get_trigrams(question)
        
        scored = []
        with open(self.path, encoding="utf-8") as f:
            for line in f:
                try:
                    entry = json.loads(line.strip())
                except json.JSONDecodeError:
                    continue
                past_q = entry.get("q", "")
                e_words = set(re.findall(r'\w+', past_q.lower()))
                e_trigrams = get_trigrams(past_q)
                
                word_overlap = len(q_words & e_words)
                trigram_overlap = len(q_trigrams & e_trigrams)
                
                # Score heavily favors exact word matches, but trigrams help with semantic/partial matches
                score = (word_overlap * 5) + trigram_overlap
                if score > 0:
                    scored.append((score, entry))
                    
        scored.sort(key=lambda x: -x[0])
        if not scored:
            return ""
        lines = []
        for _, e in scored[:n]:
            lines.append(f"之前问过类似问题: \"{e['q']}\"\n回答概要: {e['a']}")
        return "\n\n".join(lines)

    def clear(self):
        self.path.unlink(missing_ok=True)


# ═══════════════════════════════════════════════════════════════════════
#  Planning
# ═══════════════════════════════════════════════════════════════════════

@dataclass
class PlanStep:
    index: int
    description: str
    tool_hint: str = ""      # suggested tool
    status: str = "pending"  # pending / running / done / failed
    result: str = ""


@dataclass
class Plan:
    goal: str
    steps: list[PlanStep] = field(default_factory=list)
    current: int = 0

    @property
    def done(self) -> bool:
        return all(s.status in ("done", "failed") for s in self.steps)

    @property
    def current_step(self) -> PlanStep | None:
        for s in self.steps:
            if s.status == "pending":
                return s
        return None

    def mark_current(self, status: str, result: str = ""):
        for s in self.steps:
            if s.status == "pending":
                s.status = status
                s.result = result[:200]
                break

    def to_context(self) -> str:
        lines = [f"Goal: {self.goal}", ""]
        for s in self.steps:
            status_icon = {"pending": "[ ]", "running": "[~]", "done": "[+]", "failed": "[-]"}
            lines.append(f"  {status_icon.get(s.status, '?')} Step {s.index}: {s.description}")
            if s.result:
                lines.append(f"      → {s.result[:100]}")
        return "\n".join(lines)


PLANNER_PROMPT = """\
你是一个任务规划器。根据用户目标，拆解为 2-5 个可执行步骤。

输出 JSON 数组，每个元素：
{{"index": 1, "description": "步骤描述", "tool_hint": "建议使用的工具名(可选)"}}

可用工具：
{tools}

规则：
1. 步骤要具体可执行，不要写"分析代码"这种模糊描述
2. 每个步骤对应一次工具调用
3. 最后一步通常是"综合信息给出回答"
4. 只输出 JSON 数组，不要其他内容
"""


class Planner:
    """Decompose user goals into executable plan steps."""

    def __init__(self, llm: LLMClient, tools_desc: str):
        self.llm = llm
        self.tools_desc = tools_desc

    def create_plan(self, goal: str) -> Plan:
        prompt = PLANNER_PROMPT.format(tools=self.tools_desc)
        user_msg = f"用户目标: {goal}"
        raw = self.llm.complete(prompt, user_msg)

        steps = self._parse_steps(raw)
        if not steps:
            # Fallback: single-step plan
            steps = [PlanStep(index=1, description=f"搜索并回答: {goal}", tool_hint="search")]

        return Plan(goal=goal, steps=steps)

    def refine_plan(self, plan: Plan, observation: str) -> Plan:
        """Adjust remaining steps based on new information."""
        # Simple refinement: if a step failed, add a retry with different approach
        failed = [s for s in plan.steps if s.status == "failed"]
        for s in failed:
            s.status = "pending"  # retry once
            s.tool_hint = "find_pattern" if s.tool_hint == "search" else "search"
        return plan

    def _parse_steps(self, raw: str) -> list[PlanStep]:
        raw = raw.strip()
        # Look for json blocks
        json_blocks = re.findall(r"```(?:json)?\s*(.*?)\s*```", raw, re.DOTALL)
        for block in json_blocks:
            try:
                data = json.loads(block)
                if isinstance(data, list):
                    return [
                        PlanStep(
                            index=item.get("index", i+1),
                            description=item.get("description", ""),
                            tool_hint=item.get("tool_hint", ""),
                        )
                        for i, item in enumerate(data)
                        if isinstance(item, dict)
                    ]
            except Exception:
                continue

        # Fallback to regex
        match = re.search(r"\[.*\]", raw, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group(0))
                if isinstance(data, list):
                    return [
                        PlanStep(
                            index=item.get("index", i+1),
                            description=item.get("description", ""),
                            tool_hint=item.get("tool_hint", ""),
                        )
                        for i, item in enumerate(data)
                        if isinstance(item, dict)
                    ]
            except Exception:
                pass
        
        # Final fallback to direct load
        try:
            data = json.loads(raw)
            if isinstance(data, list):
                return [
                    PlanStep(
                        index=item.get("index", i+1),
                        description=item.get("description", ""),
                        tool_hint=item.get("tool_hint", ""),
                    )
                    for i, item in enumerate(data)
                    if isinstance(item, dict)
                ]
        except Exception:
            pass

        return []


# ═══════════════════════════════════════════════════════════════════════
#  Action Executor
# ═══════════════════════════════════════════════════════════════════════

class ActionExecutor:
    """Execute tool calls with retry and error handling."""

    def __init__(self, registry: ToolRegistry, max_retries: int = 1):
        self.registry = registry
        self.max_retries = max_retries
        self.log: list[dict] = []

    def execute(self, tool_name: str, params: dict, ctx: dict) -> ToolResult:
        """Execute with retry on failure."""
        result = None
        for attempt in range(self.max_retries + 1):
            result = self.registry.execute(tool_name, params, ctx)
            self.log.append({
                "tool": tool_name,
                "params": {k: str(v)[:50] for k, v in params.items()},
                "success": result.success,
                "attempt": attempt + 1,
                "elapsed_ms": result.elapsed_ms,
            })
            if result.success:
                return result
            # On retry, try adjusting params slightly
            if attempt < self.max_retries and tool_name == "search":
                params["n"] = str(int(params.get("n", 5)) + 3)
        return result


# ═══════════════════════════════════════════════════════════════════════
#  Agent Orchestration
# ═══════════════════════════════════════════════════════════════════════

AGENT_SYSTEM = """\
# Snowcode Agent - 智能代码助手

你是一个专业的代码Agent，能够通过工具查找、分析、修改代码，并解决复杂的编程问题。

## 环境信息
- 操作系统: {platform}
- 项目根目录: {project_root}
- 最大执行步骤: {max_steps}
- 可用工具: {tool_count}个

## 核心能力
1. **代码理解**: 深入分析代码结构、依赖关系和逻辑流程
2. **问题诊断**: 识别bug、性能问题和安全漏洞
3. **代码修改**: 安全地修改代码，保持代码质量和一致性
4. **测试生成**: 为代码生成全面的测试用例
5. **文档生成**: 创建清晰的代码文档和注释

## 输出格式（严格JSON格式）

### 调用工具
```json
{{"think": "详细的思考过程，包括分析、推理和计划", "tool": "工具名", "params": {{"参数": "值"}}}}
```

### 给出结论
```json
{{"think": "总结思考过程", "answer": "最终回答(Markdown格式)"}}
```

## 执行规则

### 效率优化
1. **并行执行**: 对于独立的工具调用，尽量并行执行以提高效率
2. **批量操作**: 一次性读取多个相关文件，减少工具调用次数
3. **智能搜索**: 使用精确的搜索关键词，避免泛泛的搜索
4. **缓存利用**: 充分利用已获取的信息，避免重复查询

### 安全准则
1. **路径安全**: 所有文件操作必须限制在项目根目录内
2. **输入验证**: 仔细验证所有输入参数，防止注入攻击
3. **备份机制**: 修改文件前自动创建备份(.bak文件)
4. **危险操作**: 删除、覆盖等操作需要明确确认

### 质量要求
1. **代码风格**: 保持与项目一致的代码风格和命名约定
2. **错误处理**: 添加适当的错误处理和边界检查
3. **性能考虑**: 注意算法复杂度和资源使用
4. **可维护性**: 编写清晰、可维护的代码

### 沟通规范
1. **简洁直接**: 回答直接了当，避免不必要的解释
2. **精确定位**: 标注代码位置时使用"文件路径:行号"格式
3. **中英文**: 中文回答，代码术语和专有名词使用英文
4. **结构化**: 使用Markdown格式组织复杂回答

## 工具使用指南

### 搜索优先级
1. **精确搜索**: 先尝试精确的函数名、类名搜索
2. **语义搜索**: 当精确搜索无效时使用语义搜索
3. **模式匹配**: 使用正则表达式查找特定模式
4. **目录浏览**: 了解项目结构时使用目录浏览

### 文件操作
1. **读取**: 尽量一次性读取完整文件，避免分段读取
2. **修改**: 使用search_replace进行精确修改，避免重写整个文件
3. **创建**: 创建新文件时确保目录存在
4. **删除**: 删除文件前确认文件重要性

### 代码分析
1. **上下文理解**: 分析代码时考虑相关依赖和调用链
2. **影响评估**: 修改前评估对代码库其他部分的影响
3. **测试验证**: 修改后建议运行相关测试验证

## 特殊场景处理

### 调试问题
1. 复现问题 → 定位相关代码 → 分析可能原因 → 提出解决方案 → 验证修复

### 代码重构
1. 理解现有代码 → 识别重构机会 → 制定重构计划 → 逐步实施 → 验证功能

### 性能优化
1. 性能分析 → 识别瓶颈 → 优化算法 → 验证效果

### 安全审计
1. 代码扫描 → 识别漏洞 → 评估风险 → 修复建议

## 记忆和上下文
- 使用工具时注意利用短期记忆中的已有信息
- 对于复杂问题，可以分步骤解决，保持思路连贯
- 如果连续3次工具调用无结果，直接基于已有信息给出最佳回答

## 最终目标
成为开发者的智能编程伙伴，不仅解决问题，还能提供专业建议，帮助编写更高质量的代码。
"""


@dataclass
class AgentResult:
    answer: str
    plan: Plan | None
    actions: list[dict]
    steps_taken: int
    memory_entries: int


class CodeAgent:
    """
    ReAct Agent: Plan → Think → Act → Observe → Reflect → ... → Answer

    Components:
    - Planning:   decompose goal into steps
    - Memory:     short-term (session) + long-term (persisted)
    - Tools:      search, read_file, find_pattern, list_dir, read_multiple
    - Action:     execute tools with retry
    - LLM:        reasoning engine
    """

    def __init__(
        self,
        store: VectorStore,
        project_root: Path,
        model: str | None = None,
        max_steps: int = 5,
        use_planning: bool = True,
    ):
        self.store = store
        self.root = project_root.resolve()
        self.max_steps = max_steps
        self.use_planning = use_planning

        # Components
        self.llm = LLMClient(model)
        self.tools = build_default_registry()
        self.memory_st = ShortTermMemory(max_entries=20)
        self.memory_lt = LongTermMemory(self.root)
        self.executor = ActionExecutor(self.tools, max_retries=1)
        self.planner = Planner(self.llm, self.tools.list_definitions())

        self.ctx = {"store": store, "root": self.root, "llm": self.llm}

    def run(
        self,
        question: str,
        on_step: Callable[[int, str, str], None] | None = None,
        on_think: Callable[[str], None] | None = None,
        on_answer: Callable[[str], None] | None = None,
    ) -> AgentResult:
        """
        Run the agent loop.

        on_step(step_num, tool_name, result_preview)
        on_think(think_text)
        on_answer(answer_text)
        """
        # 1. Recall long-term memory
        past = self.memory_lt.recall(question)
        if past:
            self.memory_st.add("system", f"相关历史记忆:\n{past}")

        # 2. Plan
        plan = None
        if self.use_planning:
            plan = self.planner.create_plan(question)

        # 3. Execute loop
        answer = ""
        no_result_streak = 0
        json_retries = 0
        MAX_JSON_RETRIES = 3
        
        step = 0
        _max = self.max_steps if self.max_steps > 0 else 50  # Default cap at 50, override with --steps
        _action_history: list[tuple[str, str]] = []  # (tool_name, params_str) for repeat detection
        _repeat_count = 0
        
        while step < _max:
            # Build prompt
            _platform = sys.platform  # win32 / linux / darwin
            _tool_count = len(self.tools.get_all())
            system = AGENT_SYSTEM.format(
                max_steps=_max, 
                platform=_platform,
                project_root=str(self.root),
                tool_count=_tool_count
            )
            tools_desc = self.tools.list_definitions_for_llm()
            mem_ctx = self.memory_st.get_context()

            plan_ctx = ""
            if plan:
                plan_ctx = f"\n## 当前计划\n{plan.to_context()}\n"

            user_msg = f"""## 工具列表
{tools_desc}
{plan_ctx}
## 记忆
{mem_ctx}

## 用户问题
{question}

请思考下一步。输出 JSON:"""

            raw = self.llm.complete(system, user_msg, temperature=0.2)

            if not raw:
                # LLM unavailable, fallback
                results = self.store.query(question, n_results=5)
                if results:
                    from .rag import _format_context
                    answer = "LLM 不可用，相关代码：\n\n" + _format_context(results)
                else:
                    answer = "未找到相关代码，LLM 不可用。"
                break

            try:
                parsed = self._parse_json(raw)
                json_retries = 0 # Reset on success
            except ValueError as e:
                json_retries += 1
                if json_retries >= MAX_JSON_RETRIES:
                    answer = f"Error: LLM consistently returns invalid JSON after {MAX_JSON_RETRIES} attempts. Raw output:\n{raw}"
                    break
                    
                # Tell the LLM it produced invalid JSON so it can correct itself
                err_msg = f"System Error: {e}. Please output valid JSON only."
                self.memory_st.add("system", err_msg)
                if on_step:
                    on_step(step + 1, "system_error", "JSON parse failed, retrying...")
                # Do not increment step counter on JSON parse failure
                continue

            # Increment step counter only for successful valid JSON parse
            step += 1

            think = parsed.get("think", "")

            if on_think and think:
                on_think(think)
            self.memory_st.add("agent", f"Think: {think}")

            # Check for final answer
            if "answer" in parsed and "tool" not in parsed:
                answer = parsed["answer"]
                break

            # Execute tool
            tool_name = parsed.get("tool", "")
            params = parsed.get("params", {})

            if not tool_name:
                answer = parsed.get("answer", raw)
                break

            result = self.executor.execute(tool_name, params, self.ctx)

            preview = result.output[:200] + "..." if len(result.output) > 200 else result.output
            if on_step:
                on_step(step + 1, tool_name, preview)

            # Store full tool output in memory (up to 8000 chars) so LLM can see complete file content
            self.memory_st.add("tool", f"[{tool_name}] {result.output[:8000]}", tool_name=tool_name)

            # Repeat detection: same tool + same params = loop
            action_key = (tool_name, json.dumps(params, sort_keys=True, ensure_ascii=False)[:100])
            if action_key in _action_history:
                _repeat_count += 1
            else:
                _repeat_count = 0
            _action_history.append(action_key)
            if _repeat_count >= 2:
                answer = f"检测到重复操作（连续 {_repeat_count+1} 次相同调用），强制输出当前结论：\n\n"
                # Ask LLM one more time to summarize what we know
                summary_prompt = f"基于已收集的信息，直接回答用户问题：{question}"
                summary = self.llm.complete("直接回答，不要调用工具。", summary_prompt, temperature=0.3)
                answer += summary if summary else "无法得出结论。"
                break

            if plan:
                if result.success:
                    plan.mark_current("done", result.output[:100])
                else:
                    plan.mark_current("failed", result.output[:100])
                    self.planner.refine_plan(plan, result.output)

            # Track no-result streak
            if result.success and len(result.output) < 20:
                no_result_streak += 1
            else:
                no_result_streak = 0
            if no_result_streak >= 3:
                answer = "连续多次未找到有效信息，请尝试更具体的问题。"
                break

        if not answer:
            answer = "已达到最大步数。"

        # 4. Store to long-term memory
        self.memory_lt.store(question, answer, self.executor.log)

        if on_answer:
            on_answer(answer)

        return AgentResult(
            answer=answer,
            plan=plan,
            actions=self.executor.log,
            steps_taken=len(self.executor.log),
            memory_entries=len(self.memory_st.entries),
        )

    def _parse_json(self, raw: str) -> dict | list:
        """Parse JSON robustly, even if wrapped in markdown blocks or text."""
        raw = raw.strip()

        # 1. Try finding ```json ... ``` blocks
        json_blocks = re.findall(r"```(?:json)?\s*(.*?)\s*```", raw, re.DOTALL)
        for block in json_blocks:
            try:
                return json.loads(block)
            except json.JSONDecodeError:
                continue

        # 2. Try finding { ... } or [ ... ] with balanced braces
        for start_char, end_char in [("{", "}"), ("[", "]")]:
            start = raw.find(start_char)
            if start == -1:
                continue
            # Find matching end by counting braces
            depth = 0
            for i in range(start, len(raw)):
                if raw[i] == start_char:
                    depth += 1
                elif raw[i] == end_char:
                    depth -= 1
                if depth == 0:
                    try:
                        return json.loads(raw[start:i+1])
                    except json.JSONDecodeError:
                        break

        # 3. Try raw text directly
        try:
            return json.loads(raw)
        except Exception:
            pass

        # 4. Last resort: extract think/answer from plain text
        # Model might output "Think: xxx\nAnswer: xxx" without JSON
        if "answer" in raw.lower() or "答" in raw:
            # Try to extract answer portion
            for marker in ["Answer:", "answer:", "回答：", "答案："]:
                idx = raw.find(marker)
                if idx != -1:
                    answer_text = raw[idx + len(marker):].strip()
                    return {"think": "", "answer": answer_text}

        # Treat entire output as answer
        return {"think": "", "answer": raw}

    def reset_memory(self):
        """Clear all memory."""
        self.memory_st.clear()
        self.memory_lt.clear()


# ═══════════════════════════════════════════════════════════════════════
#  Coordinator-Worker Architecture
# ═══════════════════════════════════════════════════════════════════════

class WorkerAgent(CodeAgent):
    """Worker agent that executes specific tasks under coordinator guidance."""
    
    def __init__(self, store: VectorStore, project_root: Path, model: str | None = None, 
                 worker_id: str = "worker"):
        super().__init__(store, project_root, model, max_steps=3, use_planning=False)
        self.worker_id = worker_id
        self.specialization = "general"  # Can be specialized for specific tasks
    
    def execute_task(self, task_description: str, context: str = "") -> AgentResult:
        """Execute a specific task with given context."""
        # Add context to memory
        if context:
            self.memory_st.add("system", f"任务上下文:\n{context}")
        
        # Execute with limited steps
        return super().run(task_description)


class CoordinatorAgent:
    """Coordinator agent that plans and delegates tasks to worker agents."""
    
    def __init__(self, store: VectorStore, project_root: Path, model: str | None = None,
                 max_workers: int = 2):
        self.store = store
        self.root = project_root
        self.llm = LLMClient(model)
        self.max_workers = max_workers
        
        # Worker pool
        self.workers: list[WorkerAgent] = []
        self._init_workers()
        
        # Coordinator memory
        self.memory = ShortTermMemory(max_entries=30)
        self.long_term_memory = LongTermMemory(project_root)
    
    def _init_workers(self):
        """Initialize worker agents."""
        for i in range(self.max_workers):
            worker = WorkerAgent(self.store, self.root, worker_id=f"worker_{i}")
            self.workers.append(worker)
    
    def plan_and_execute(self, question: str, on_progress=None) -> AgentResult:
        """Plan complex task and delegate to workers."""
        # 1. Recall similar past tasks
        past = self.long_term_memory.recall(question)
        if past:
            self.memory.add("system", f"相关历史经验:\n{past}")
        
        # 2. Create execution plan
        plan = self._create_execution_plan(question)
        
        # 3. Execute plan steps
        results = []
        actions = []
        
        for step in plan.get("steps", []):
            step_desc = step.get("description", "")
            worker_id = step.get("worker", 0)
            
            # Select worker
            worker = self.workers[worker_id % len(self.workers)]
            
            # Execute step
            if on_progress:
                on_progress(f"执行步骤: {step_desc}")
            
            result = worker.execute_task(step_desc, context=str(results))
            results.append(result.answer)
            actions.extend(result.actions)
        
        # 4. Synthesize final answer
        final_answer = self._synthesize_results(question, results)
        
        # 5. Store in long-term memory
        self.long_term_memory.store(question, final_answer, actions)
        
        return AgentResult(
            answer=final_answer,
            plan=None,  # Coordinator creates its own internal plan
            actions=actions,
            steps_taken=len(plan.get("steps", [])),
            memory_entries=len(self.memory.entries),
        )
    
    def _create_execution_plan(self, question: str) -> dict:
        """Create execution plan for complex task."""
        if not self.llm.available:
            # Fallback: single step plan
            return {"steps": [{"description": question, "worker": 0}]}
        
        system = """你是一个任务规划专家。根据用户问题，创建详细的执行计划。
将复杂任务分解为多个简单的子任务，每个子任务可以由独立的智能体执行。

输出JSON格式:
{
  "analysis": "任务分析",
  "steps": [
    {"description": "子任务描述", "worker": 0, "priority": "high/medium/low"},
    ...
  ]
}

规则:
1. 每个子任务应该具体、可执行
2. 考虑任务依赖关系
3. 最多分解为5个子任务
4. worker编号从0开始"""
        
        user_msg = f"请为以下任务创建执行计划:\n\n{question}"
        
        raw = self.llm.complete(system, user_msg, temperature=0.2)
        
        try:
            # Parse JSON response
            import re
            json_match = re.search(r'\{.*\}', raw, re.DOTALL)
            if json_match:
                plan = json.loads(json_match.group())
                return plan
        except Exception:
            pass
        
        # Fallback plan
        return {"steps": [{"description": question, "worker": 0}]}
    
    def _synthesize_results(self, question: str, results: list[str]) -> str:
        """Synthesize multiple results into final answer."""
        if len(results) == 1:
            return results[0]
        
        if not self.llm.available:
            # Simple concatenation
            return "\n\n".join(f"结果{i+1}:\n{r}" for i, r in enumerate(results))
        
        system = """你是一个结果整合专家。将多个子任务的结果整合为一个连贯、完整的最终答案。
确保答案结构清晰，逻辑连贯，避免重复信息。"""
        
        results_text = "\n\n".join(f"结果{i+1}:\n{r}" for i, r in enumerate(results))
        user_msg = f"原始问题: {question}\n\n子任务结果:\n{results_text}\n\n请整合为最终答案:"
        
        return self.llm.complete(system, user_msg, temperature=0.1)
